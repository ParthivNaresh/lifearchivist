import asyncio
import csv
import io
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import List, Tuple

import aiofiles
import chardet
import pytesseract
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook
from openpyxl.cell import Cell
from openpyxl.utils.exceptions import InvalidFileException
from PIL import Image, ImageEnhance, ImageFilter
from pypdf import PdfReader


def _get_extraction_method(mime_type: str) -> str:
    """Get extraction method name based on mime type."""
    if mime_type == "text/csv":
        return "csv"
    elif mime_type.startswith("text/"):
        return "text_file"
    elif mime_type == "application/pdf":
        return "pypdf"
    elif (
        mime_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        return "python_docx"
    elif mime_type in [
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",  # .xlsx
        "application/vnd.ms-excel",  # .xls
        "application/excel",  # Alternative MIME type
        "application/x-excel",  # Alternative MIME type
        "application/x-msexcel",  # Alternative MIME type
    ]:
        return "openpyxl"
    elif mime_type.startswith("image/"):
        return "tesseract"
    else:
        raise ValueError(f"Mime type not supported: {mime_type}")


async def _extract_text_file(file_path: Path) -> str:
    """Extract text from plain text files."""
    async with aiofiles.open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return await f.read()


def extract_table_text(table: Table) -> str:
    """Extract text from table with structure preservation."""
    table_text = []
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            cell_text = cell.text.strip()
            if cell_text:
                row_text.append(cell_text)
        if row_text:
            table_text.append(" | ".join(row_text))
    return "\n".join(table_text)


def extract_paragraph_text(paragraph: Paragraph) -> str:
    """Extract text from paragraph with run-level formatting awareness."""
    if not paragraph.text.strip():
        return ""

    paragraph_text = paragraph.text.strip()

    # Preserve important formatting indicators
    if paragraph.style and paragraph.style.name:
        style_name = paragraph.style.name.lower()
        if "heading" in style_name:
            # Mark headings for better semantic understanding
            paragraph_text = f"[HEADING] {paragraph_text}"
        elif "title" in style_name:
            paragraph_text = f"[TITLE] {paragraph_text}"

    return paragraph_text


async def _extract_docx_text(file_path: Path) -> str:
    """Extract text from Word documents using python-docx with comprehensive content extraction."""
    try:
        doc = Document(str(file_path))
        extracted_content = []

        # Process document body elements in order to preserve structure
        for element in doc.element.body:
            if isinstance(element, CT_P):
                # Paragraph element
                paragraph = Paragraph(element, doc)
                paragraph_text = extract_paragraph_text(paragraph)
                if paragraph_text:
                    extracted_content.append(paragraph_text)

            elif isinstance(element, CT_Tbl):
                # Table element
                table = Table(element, doc)
                table_text = extract_table_text(table)
                if table_text:
                    extracted_content.append(f"[TABLE]\n{table_text}")

        # Process headers and footers for comprehensive extraction
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    header_text = extract_paragraph_text(paragraph)
                    if header_text:
                        extracted_content.insert(0, f"[HEADER] {header_text}")

            if section.footer:
                for paragraph in section.footer.paragraphs:
                    footer_text = extract_paragraph_text(paragraph)
                    if footer_text:
                        extracted_content.append(f"[FOOTER] {footer_text}")

        # Join all content with proper spacing
        full_text = "\n\n".join(extracted_content)

        # Clean up excessive whitespace while preserving structure
        import re

        full_text = re.sub(r"\n{3,}", "\n\n", full_text)
        full_text = re.sub(r"[ \t]+", " ", full_text)
        return full_text.strip()

    except Exception as e:
        raise ValueError(
            f"Error extracting Word document text from {file_path}: {e}"
        ) from None


async def _pdf_needs_ocr(
    pdf_reader: PdfReader, min_chars_per_page: int = 200, min_unique_words: int = 20
) -> bool:
    """
    Determine if a PDF needs OCR based on text quality.

    A PDF needs OCR if:
    1. No text layer at all (no fonts in resources)
    2. Text exists but is minimal/repetitive (e.g., watermarks only)

    Args:
        pdf_reader: PyPDF PdfReader instance
        min_chars_per_page: Minimum characters per page to consider text sufficient
        min_unique_words: Minimum unique words to consider text meaningful

    Returns:
        True if OCR is needed, False otherwise
    """
    # Sample first 3 pages (or all if fewer)
    pages_to_check = min(3, len(pdf_reader.pages))

    if pages_to_check == 0:
        return False  # Empty PDF

    total_text = []
    pages_with_fonts = 0
    pages_with_sufficient_text = 0

    for i in range(pages_to_check):
        page = pdf_reader.pages[i]

        # Check for fonts in resources (indicates text layer)
        resources = page.get("/Resources", {})
        if "/Font" in resources:
            pages_with_fonts += 1

        # Extract and analyze text
        text = page.extract_text() or ""
        text_clean = text.strip()
        total_text.append(text_clean)

        # Count unique words to detect repetitive content
        words = re.findall(r"\b[a-zA-Z]+\b", text_clean.lower())
        unique_words = set(words)

        # Check if text is sufficient and diverse
        if (
            len(text_clean) >= min_chars_per_page
            and len(unique_words) >= min_unique_words
        ):
            pages_with_sufficient_text += 1

    # Analyze combined text from sampled pages
    combined_text = " ".join(total_text)
    all_words = re.findall(r"\b[a-zA-Z]+\b", combined_text.lower())
    unique_words_total = set(all_words)

    # Decision logic
    if pages_with_fonts == 0:
        # No text layer at all
        logging.info("PDF needs OCR: No text layer found (no fonts in resources)")
        return True

    if pages_with_sufficient_text == 0 and len(unique_words_total) < min_unique_words:
        # Has fonts but insufficient meaningful text (likely just watermarks)
        logging.info(
            f"PDF needs OCR: Insufficient meaningful text (only {len(unique_words_total)} unique words)"
        )
        return True

    return False


async def _extract_pdf_with_ocr(file_path: Path) -> str:
    """
    Extract text from PDF using OCR by converting pages to images.

    Args:
        file_path: Path to the PDF file

    Returns:
        Extracted text from all pages
    """
    try:
        import tempfile

        from pdf2image import convert_from_path

        logging.info(f"Starting OCR extraction for PDF: {file_path.name}")

        # Convert PDF pages to images
        # Use lower DPI for faster processing (200 is usually sufficient for text)
        with tempfile.TemporaryDirectory() as temp_dir:
            images = convert_from_path(
                str(file_path),
                dpi=200,  # Balance between quality and speed
                output_folder=temp_dir,
                fmt="PNG",
                thread_count=4,  # Use multiple threads for faster conversion
                use_pdftocairo=True,  # More reliable than pdftoppm
            )

            logging.info(f"Converted {len(images)} pages to images for OCR")

            all_text: List[str] = []

            for page_num, image in enumerate(images, 1):
                logging.info(f"Processing page {page_num}/{len(images)} with OCR")

                # Preprocess image for better OCR
                processed_image = await _preprocess_image_for_ocr(image)

                # Run OCR
                try:
                    # Configure Tesseract for better accuracy
                    custom_config = r"--oem 3 --psm 3"  # OEM 3 = Default, PSM 3 = Fully automatic page segmentation

                    text = await asyncio.to_thread(
                        pytesseract.image_to_string,
                        processed_image,
                        config=custom_config,
                    )

                    if text.strip():
                        # Add page separator for multi-page documents
                        if all_text:
                            all_text.append(f"\n--- Page {page_num} ---\n")
                        all_text.append(text.strip())

                except Exception as ocr_error:
                    logging.warning(f"OCR failed for page {page_num}: {ocr_error}")
                    # Continue with other pages

            if all_text:
                logging.info(
                    f"OCR extraction complete: {len(all_text)} pages with text"
                )
                return "\n".join(all_text)
            else:
                logging.warning("No text extracted via OCR from any page")
                return ""

    except ImportError as e:
        raise ValueError(
            "pdf2image is not installed. Please install it using: "
            "pip install pdf2image and install poppler-utils (brew install poppler on macOS)"
        ) from e
    except Exception as e:
        logging.error(f"Error during PDF OCR extraction: {e}")
        raise ValueError(f"Error extracting PDF text via OCR: {e}") from None


async def _extract_pdf_text(file_path: Path) -> str:
    """
    Extract text from PDF files with intelligent OCR fallback.

    First attempts standard text extraction using PyPDF.
    If the PDF has no text layer or only minimal/repetitive text (e.g., watermarks),
    falls back to OCR extraction.

    Args:
        file_path: Path to the PDF file

    Returns:
        Extracted text content
    """
    try:
        # First attempt: Standard text extraction
        text_content = []
        needs_ocr = False

        with open(file_path, "rb") as file:
            pdf_reader = PdfReader(file)

            # Check if OCR is needed
            needs_ocr = await _pdf_needs_ocr(pdf_reader)

            if not needs_ocr:
                # Standard extraction is sufficient
                logging.info(
                    f"Using standard text extraction for PDF: {file_path.name}"
                )
                for page in pdf_reader.pages:
                    text_content.append(page.extract_text())

                extracted_text = "\n".join(text_content)

                # Final sanity check - if we got very little text from a multi-page PDF
                if len(pdf_reader.pages) > 5 and len(extracted_text.strip()) < 500:
                    logging.warning(
                        f"Suspiciously little text ({len(extracted_text)} chars) from {len(pdf_reader.pages)} pages, trying OCR"
                    )
                    needs_ocr = True

        # If standard extraction failed or was insufficient, use OCR
        if needs_ocr:
            logging.info(f"PDF requires OCR extraction: {file_path.name}")
            extracted_text = await _extract_pdf_with_ocr(file_path)

        return extracted_text

    except Exception as e:
        logging.error(f"Error extracting PDF text from {file_path}: {e}")
        # Try OCR as last resort
        try:
            logging.info(
                f"Standard extraction failed, attempting OCR fallback for: {file_path.name}"
            )
            return await _extract_pdf_with_ocr(file_path)
        except Exception as ocr_error:
            raise ValueError(
                f"Both standard and OCR extraction failed: {ocr_error}"
            ) from None


def _format_cell_value(cell: Cell) -> str:
    """
    Format cell value for text extraction, handling various data types.

    Args:
        cell: The openpyxl cell object

    Returns:
        Formatted string representation of the cell value
    """
    from lifearchivist.tools.extract.utils import CellValueFormatter

    value = cell.value

    if value is None:
        return ""

    if isinstance(value, datetime):
        return CellValueFormatter.format_datetime(value)

    if isinstance(value, bool):
        return CellValueFormatter.format_boolean(value)

    if isinstance(value, (int, float)):
        return CellValueFormatter.format_numeric_value(cell, value)

    if hasattr(cell, "data_type") and cell.data_type == "f":
        return CellValueFormatter.format_formula_result(cell, value)

    return str(value).strip()


def _detect_header_row(sheet, max_rows: int = 10) -> Tuple[int, List[str]]:
    """
    Intelligently detect header row in a spreadsheet.

    Args:
        sheet: The openpyxl worksheet object
        max_rows: Maximum number of rows to check for headers

    Returns:
        Tuple of (header_row_index, list_of_headers)
    """
    for row_idx in range(1, min(max_rows + 1, sheet.max_row + 1)):
        row_values = []
        non_empty_count = 0

        for col_idx in range(1, sheet.max_column + 1):
            cell = sheet.cell(row=row_idx, column=col_idx)
            value = _format_cell_value(cell)
            row_values.append(value)
            if value:
                non_empty_count += 1

        # Check if this row could be a header
        # Headers typically have mostly non-empty text values
        if non_empty_count >= sheet.max_column * 0.5:  # At least 50% filled
            # Check if values look like headers (text, not numbers)
            text_count = sum(
                1 for v in row_values if v and not re.match(r"^-?\d+\.?\d*$", v)
            )
            if text_count >= non_empty_count * 0.7:  # At least 70% text
                return row_idx, row_values

    # No clear header found, use first row if it has content
    first_row = []
    for col_idx in range(1, sheet.max_column + 1):
        cell = sheet.cell(row=1, column=col_idx)
        first_row.append(_format_cell_value(cell))

    if any(first_row):
        return 1, first_row

    return 0, []


async def _extract_excel_text(file_path: Path) -> str:
    """
    Extract text from Excel files with comprehensive data handling.

    Handles:
    - Multiple sheets
    - Various data types (dates, numbers, currency, percentages)
    - Formulas (extracts calculated values)
    - Merged cells
    - Empty cells and sparse data
    - Large spreadsheets efficiently
    """
    from lifearchivist.tools.extract.utils import ExcelSheetProcessor

    try:
        workbook = load_workbook(
            str(file_path),
            read_only=True,
            data_only=True,
            keep_links=False,
        )

        extracted_content = []

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]

            sheet_content = ExcelSheetProcessor.process_single_sheet(
                sheet, sheet_name, _format_cell_value, _detect_header_row
            )

            if sheet_content:
                extracted_content.extend(sheet_content)
                extracted_content.append("")

        workbook.close()

        full_text = "\n".join(extracted_content)
        full_text = re.sub(r"\n{4,}", "\n\n\n", full_text)

        return full_text.strip()

    except InvalidFileException as e:
        raise ValueError(f"Invalid Excel file format for {file_path}: {e}") from None
    except Exception as e:
        raise ValueError(f"Error extracting Excel text from {file_path}: {e}") from None


async def _detect_csv_encoding(file_path: Path, sample_size: int = 10000) -> str:
    """
    Detect the encoding of a CSV file.

    Args:
        file_path: Path to the CSV file
        sample_size: Number of bytes to sample for detection

    Returns:
        Detected encoding string
    """
    try:
        with open(file_path, "rb") as f:
            raw_data = f.read(sample_size)

        # Try to detect encoding
        result = chardet.detect(raw_data)
        encoding = result.get("encoding", "utf-8")
        confidence = result.get("confidence", 0)

        # Fallback for low confidence or None result
        if not encoding or confidence < 0.7:
            # Try common encodings
            for enc in ["utf-8", "latin-1", "cp1252", "iso-8859-1"]:
                try:
                    raw_data.decode(enc)
                    return enc
                except UnicodeDecodeError:
                    continue
            # Default to utf-8 with error handling
            return "utf-8"

        return encoding

    except Exception:
        return "utf-8"


def _detect_csv_delimiter(content: str, sample_lines: int = 10) -> str:
    """
    Detect the delimiter used in CSV content.

    Args:
        content: CSV file content as string
        sample_lines: Number of lines to sample

    Returns:
        Detected delimiter character
    """
    lines = content.split("\n")[:sample_lines]

    # Common delimiters to check
    delimiters = [",", ";", "\t", "|"]
    delimiter_counts = {d: 0 for d in delimiters}

    for line in lines:
        if not line.strip():
            continue
        for delimiter in delimiters:
            # Count occurrences, considering quoted fields
            count = len(list(csv.reader([line], delimiter=delimiter))[0]) - 1
            delimiter_counts[delimiter] += count

    # Return delimiter with highest count
    if delimiter_counts:
        return max(delimiter_counts, key=lambda d: delimiter_counts.get(d, 0))
    return ","  # Default to comma


async def _preprocess_image_for_ocr(image: Image.Image) -> Image.Image:
    """
    Preprocess image to improve OCR accuracy.

    Args:
        image: PIL Image object

    Returns:
        Preprocessed PIL Image object
    """
    # Convert to RGB if necessary
    if image.mode != "RGB":
        image = image.convert("RGB")

    # Resize if image is too small (OCR works better with larger images)
    width, height = image.size
    if width < 1000 or height < 1000:
        scale_factor = max(1000 / width, 1000 / height)
        new_width = int(width * scale_factor)
        new_height = int(height * scale_factor)
        image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)

    # Convert to grayscale for better OCR
    image = image.convert("L")

    # Enhance contrast
    contrast_enhancer = ImageEnhance.Contrast(image)
    image = contrast_enhancer.enhance(1.5)

    # Apply slight sharpening
    image = image.filter(ImageFilter.SHARPEN)

    # Enhance brightness if image is too dark
    brightness_enhancer = ImageEnhance.Brightness(image)
    # Calculate average brightness
    pixels = list(image.getdata())
    avg_brightness = sum(pixels) / len(pixels)
    if avg_brightness < 127:  # If darker than middle gray
        image = brightness_enhancer.enhance(1.2)

    return image


async def _extract_image_text(file_path: Path) -> str:
    """
    Extract text from images using Tesseract OCR.

    Handles:
    - Various image formats (JPEG, PNG, GIF, BMP, TIFF, WebP)
    - Image preprocessing for better OCR accuracy
    - Multi-page TIFF files
    - Rotated or skewed text
    - Low quality images

    Returns:
        Extracted text only (no metadata markers)
    """
    from lifearchivist.tools.extract.utils import OCRProcessor

    try:
        image = Image.open(file_path)

        logging.info(
            f"OCR processing image: Format={image.format}, Size={image.size[0]}x{image.size[1]}, Mode={image.mode}"
        )

        num_pages = OCRProcessor.get_num_pages(image)
        all_text: List[str] = []
        total_confidence: List[int] = []

        for page_num in range(num_pages):
            text, confidences = await OCRProcessor.process_single_page(
                image, page_num, num_pages, _preprocess_image_for_ocr
            )

            if text:
                all_text.append(text)
                total_confidence.extend(confidences)

        formatted_text = OCRProcessor.format_multipage_text(all_text, num_pages)
        OCRProcessor.log_ocr_summary(num_pages, all_text, total_confidence)

        return "\n".join(formatted_text) if formatted_text else ""

    except Exception as e:
        raise ValueError(f"Error extracting text from image {file_path}: {e}") from None


async def _extract_csv_text(file_path: Path) -> str:
    """
    Extract text from CSV files with intelligent parsing.

    Handles:
    - Various encodings (UTF-8, Latin-1, Windows-1252, etc.)
    - Different delimiters (comma, semicolon, tab, pipe)
    - Quoted fields with embedded delimiters
    - Large files efficiently
    - Malformed CSV data
    - Date and number formatting
    """
    from lifearchivist.tools.extract.utils import CSVProcessor

    try:
        encoding = await _detect_csv_encoding(file_path)
        content = await CSVProcessor.read_csv_content(file_path, encoding)
        delimiter = _detect_csv_delimiter(content)

        csv_reader = csv.reader(io.StringIO(content), delimiter=delimiter)

        extracted_rows = []
        header_row = None
        row_count = 0
        max_rows = 10000

        for row_idx, row in enumerate(csv_reader):
            if not any(cell.strip() for cell in row):
                continue

            cleaned_row = CSVProcessor.process_csv_row(row)

            if row_idx == 0 and any(cleaned_row):
                header_row = cleaned_row
                extracted_rows.extend(CSVProcessor.format_header_row(cleaned_row))
            else:
                extracted_rows.append(CSVProcessor.format_data_row(cleaned_row))

            row_count += 1
            if row_count >= max_rows:
                extracted_rows.append(f"[... truncated after {max_rows} rows ...]")
                break

        summary = CSVProcessor.create_summary(
            row_count, delimiter, encoding, header_row
        )
        extracted_rows.extend(summary)

        return "\n".join(extracted_rows)

    except Exception as e:
        raise ValueError(f"Error extracting CSV text from {file_path}: {e}") from None


async def _extract_text_by_type(file_path: Path, mime_type: str) -> str:
    """Extract text based on file type."""
    try:
        # CSV files (check first as they can be text/*)
        if mime_type == "text/csv":
            return await _extract_csv_text(file_path)
        # Plain text files
        elif mime_type.startswith("text/"):
            return await _extract_text_file(file_path)
        # PDF files
        elif mime_type == "application/pdf":
            return await _extract_pdf_text(file_path)
        # Word documents
        elif (
            mime_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            return await _extract_docx_text(file_path)
        # Excel files
        elif mime_type in [
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",  # .xlsx
            "application/vnd.ms-excel",  # .xls
            "application/excel",
            "application/x-excel",
            "application/x-msexcel",
        ]:
            return await _extract_excel_text(file_path)
        # Image files
        elif mime_type.startswith("image/"):
            return await _extract_image_text(file_path)
        else:
            return ""
    except Exception as e:
        raise ValueError(f"Error extracting text from {file_path}: {e}") from None
